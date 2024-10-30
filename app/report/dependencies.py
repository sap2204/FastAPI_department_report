from collections import defaultdict
from datetime import date, datetime
from docx import Document



from app.tasks.schemas import STasksAllWorkers




def create_report(starting_date: date, 
                    end_date: date,
                    data_for_report: list[STasksAllWorkers]):
    
    # Перевод даты из формата "Год-месяц-день" в формат "День-Месяц-Год"
    formated_starting_date = starting_date.strftime("%d-%m-%Y")
    formated_end_date = end_date.strftime("%d-%m-%Y")
    
    # Группировка задач по работникам
    grouped_data = defaultdict(list)
    
    # Прохожу по списку из словарей. В каждом словаре формирую полное имя full_name
    # Формирую новый словарь. Обращаюсь по ключу (полное имя), которого нет, но defaultdict 
    # создает ключ (полное имя), а значением этого ключа является список задач и количество дней на задачу
    for item in data_for_report:
        full_name = f"{item['surname']} {item['name']} {item['middle_name']}"
        grouped_data[full_name].append(f"{item['task']} - {item['days_per_task']} дней;")

    
    # Создание нового документа
    doc = Document()

    # Заголовок в отчете и его выравнивание по левому краю по умолчанию
    doc.add_heading('МиСНК \n', level=3)

    # Добавление в отчет информации о периоде
    doc.add_paragraph(f"Период отчетности: {formated_starting_date} - {formated_end_date}\n")

    # Создаю таблицу с двумя колонками
    table = doc.add_table(rows=1, cols=2)
    table.autofit = True
    
    # Даю названия колонкам
    cells = table.rows[0].cells # получил ячейки первой строки (для заголовков)
    cells[0].text = "Исполнитель" # название 1-го столбца
    cells[1].text = "Основные задачи" # название 2-го столбца

    # Добавление записей из словаря с группированными задачами по работнику
    for full_name, tasks in grouped_data.items():
        # Создаю новую ячейку для записи работника
        row_cells = table.add_row().cells

        # Записываю работника в левую ячейку
        row_cells[0].text = full_name

        # Записываю задачи в правую ячейку
        row_cells[1].text = "\n".join(tasks) # Каждая задача записывается с новой строки

    # Сохранение отчета
    doc.save(f"C:\\Работа\\1_Отчеты отдела по месяцам\\Отчет МиСНК {formated_starting_date} - {formated_end_date}.docx")



